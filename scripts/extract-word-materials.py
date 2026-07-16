from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document


def clean_text(value: str) -> str:
    value = value.replace("\u3000", " ")
    value = re.sub(r"[ \t]+", " ", value)
    return value.strip()


def read_docx(path: Path) -> dict:
    doc = Document(str(path))
    blocks: list[dict] = []

    for paragraph in doc.paragraphs:
        text = clean_text(paragraph.text)
        if not text:
            continue
        style = paragraph.style.name if paragraph.style else ""
        blocks.append({"type": "paragraph", "style": style, "text": text})

    for table_index, table in enumerate(doc.tables, start=1):
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            blocks.append({"type": "table", "index": table_index, "rows": rows})

    return {
        "sourceFile": str(path),
        "title": path.stem,
        "blocks": blocks,
    }


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: extract-word-materials.py <input.docx> <output.json>", file=sys.stderr)
        return 2

    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    data = read_docx(source)
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"extracted {len(data['blocks'])} blocks -> {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
